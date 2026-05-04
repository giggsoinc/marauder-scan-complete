"""Build Pitch_Data_Sheet.docx — 2-page CXO pitch for PatronAI.

Run:  python3 scripts/build_pitch_docx.py
Out:  Pitch_Data_Sheet.docx (repo root)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND       = RGBColor(0x0B, 0x3D, 0x91)
BRAND_LIGHT = RGBColor(0x17, 0x69, 0xD8)
INK         = RGBColor(0x1A, 0x1A, 0x1A)
MUTED       = RGBColor(0x5A, 0x5A, 0x5A)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_run(para, text, *, bold=False, size=10, color=INK):
    r = para.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    return r


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin   = Cm(1.6)
        section.right_margin  = Cm(1.6)

    # ── Banner header ─────────────────────────────────────────
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = True
    bcell = banner.rows[0].cells[0]
    set_cell_bg(bcell, "0B3D91")
    bcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = bcell.paragraphs[0]
    add_run(p, "PATRONAI  ", bold=True, size=20, color=RGBColor(0xFF, 0xFF, 0xFF))
    add_run(p, "Find the AI you didn't know you had.\n",
            size=11, color=RGBColor(0xE5, 0xEC, 0xF5))
    add_run(p, "Shadow AI · Ghost AI · Unmanaged Models — discovered, mapped, alerted in one container.",
            size=9, color=RGBColor(0xC8, 0xD4, 0xE8))

    # ── The Problem ───────────────────────────────────────────
    h = doc.add_paragraph()
    add_run(h, "THE PROBLEM CXOs ARE LOSING SLEEP OVER", bold=True, size=11, color=BRAND)

    p = doc.add_paragraph()
    add_run(p,
        "Your developers are using AI you didn't approve. Your data is "
        "leaving your network through ChatGPT, Copilot, Claude, Cursor, and "
        "40+ AI frameworks — at commit time, in browser tabs, through MCP "
        "servers nobody catalogued. Compliance has no audit trail. Security "
        "has no inventory. Legal has questions you can't answer.",
        size=10)

    # ── The Solution ──────────────────────────────────────────
    h = doc.add_paragraph()
    add_run(h, "WHAT PATRONAI DOES (IN ONE BREATH)", bold=True, size=11, color=BRAND)

    p = doc.add_paragraph()
    add_run(p,
        "Single Docker container. Deploys in 10 minutes on a $60/month EC2. "
        "Watches three layers — network traffic, source code at commit, "
        "and developer endpoints — and tells you exactly which AI tools "
        "your team uses, which API keys are in your repos, and which "
        "models are running on your laptops. Data never leaves your cloud. "
        "Apache 2.0.",
        size=10)

    # ── Use Cases Solved (3-col table) ────────────────────────
    h = doc.add_paragraph()
    add_run(h, "USE CASES SOLVED", bold=True, size=11, color=BRAND)

    uc = doc.add_table(rows=4, cols=2)
    uc.style = "Light Grid Accent 1"

    use_cases = [
        ("Shadow AI Discovery",
         "“Which AI tools is my team using?” — Network layer matches "
         "outbound traffic to 70+ AI providers (OpenAI, Anthropic, Cohere, "
         "Mistral, Bedrock, Vertex, Azure OpenAI…). Alert in under 5 min."),
        ("Ghost AI in Code",
         "“Are there hardcoded API keys or unauthorized AI frameworks in "
         "our repos?” — Git pre-commit hook catches LangChain, CrewAI, "
         "AutoGen, MCP server configs, and OpenAI/Anthropic key literals "
         "before they merge."),
        ("Endpoint AI Inventory",
         "“What AI is running on developer laptops?” — Per-device agent "
         "(macOS · Linux · Windows) scans pip/npm/brew packages, running "
         "processes, browser history, IDE plugins, Docker containers — "
         "every 30 min."),
        ("On-Prem AI Chat over Findings",
         "“Who is the riskiest user this quarter? Show me shadow AI "
         "by provider.” — Built-in LLM (LFM2.5 via llama.cpp) answers "
         "natural-language questions over your own findings. Zero cloud "
         "AI calls."),
    ]
    for i, (title, body) in enumerate(use_cases):
        c0, c1 = uc.rows[i].cells
        c0.width = Cm(4.5)
        c1.width = Cm(13.0)
        p0 = c0.paragraphs[0]
        add_run(p0, title, bold=True, size=10, color=BRAND_LIGHT)
        p1 = c1.paragraphs[0]
        add_run(p1, body, size=9.5)

    # ── Proof / Differentiators ───────────────────────────────
    h = doc.add_paragraph()
    add_run(h, "WHY PATRONAI VS BUILD-IT-YOURSELF / DLP / CASB", bold=True, size=11, color=BRAND)

    bullets = [
        ("Three layers in one product.", " Network + code + endpoint. DLP gives you network only; CASB gives you SaaS only; secret scanners give you code only."),
        ("Zero data egress.", " Findings + chat both stay in your AWS account. The on-prem LLM means even AI Q&A never leaves the perimeter."),
        ("OCSF-native.", " Every event normalised to industry-standard OCSF — drops straight into Splunk, Sentinel, Chronicle."),
        ("Scales to millions of findings.", " Hourly S3 rollups make chat sub-second at multi-month volume."),
        ("Apache 2.0.", " No procurement cycle. Try it today. Pay only for commercial support if/when you want it."),
    ]
    for label, rest in bullets:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, label, bold=True, size=10)
        add_run(p, rest, size=10)

    # ── The Ask ───────────────────────────────────────────────
    ask_table = doc.add_table(rows=1, cols=1)
    acell = ask_table.rows[0].cells[0]
    set_cell_bg(acell, "F1F5FB")
    p = acell.paragraphs[0]
    add_run(p, "THE ASK\n", bold=True, size=11, color=BRAND)
    add_run(p,
        "30 minutes — we'll deploy PatronAI in your sandbox account "
        "and show you the shadow AI living on your network and in your "
        "repos. By the end of the call you'll have a list of every AI "
        "tool your team is using, ranked by risk. No commitment.",
        size=10)

    # ── Footer ────────────────────────────────────────────────
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(f, "Giggso Inc · Ravi Venugopal, Founder · ", size=9, color=MUTED)
    add_run(f, "rv@giggso.com", size=9, color=BRAND_LIGHT)
    add_run(f, " · github.com/giggsoinc/patronai · Apache 2.0",
            size=9, color=MUTED)

    out = "Pitch_Data_Sheet.docx"
    doc.save(out)
    print(f"✓ Wrote {out}")


if __name__ == "__main__":
    main()
